import streamlit as st
import google.generativeai as genai
import PyPDF2
import os
import glob
import requests
import re
import json
from docx import Document
from io import BytesIO
from bs4 import BeautifulSoup

st.set_page_config(page_title="audskal의 학교생활기록부 분석", layout="wide")
st.title("🏫 객관적이고 체계적인 학생부 분석")
st.markdown("API 키에 맞는 최적의 AI 모델을 자동으로 찾아내어 생기부를 체계적으로 분석합니다.")

@st.cache_data(show_spinner=False)
def load_reference_pdfs(pdf_list):
    text = ""
    for pdf_file in pdf_list:
        with open(pdf_file, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
    return text

with st.sidebar:
    st.header("🔑 기본 설정")
    
    # --- 💡 [신규 기능] API 제공자 선택 옵션 추가 ---
    api_provider = st.radio("🤖 API 제공자 선택", ["Google AI Studio", "OpenRouter"])
    api_key = st.text_input("🔑 API 키를 입력하세요", type="password")
    
    if api_provider == "OpenRouter":
        or_model = st.selectbox("사용할 모델 선택", ["google/gemini-2.5-flash", "anthropic/claude-3.5-sonnet", "openai/gpt-4o"])
        st.markdown("[🔗 OpenRouter 무료 API 키 발급](https://openrouter.ai/keys)")
    else:
        st.markdown("[👉 Google AI Studio 무료 API 키 발급](https://aistudio.google.com/app/apikey)")
    
    st.markdown("---")
    st.subheader("📚 내장된 기본 평가 기준 파일")
    pdf_files = glob.glob("*.pdf")
    if pdf_files:
        for f in pdf_files:
            st.write(f"- {f}")
    else:
        st.error("폴더에 기준 PDF 파일이 없습니다!")

col1, col2 = st.columns([1, 1])

with col1:
    st.subheader("1. 학교생활기록부 데이터 입력")
    st.info("💡 나이스(NEIS) 원본 PDF는 보안상 안 읽히는 경우가 많습니다. 가급적 아래 빈칸에 내용을 직접 긁어서 붙여넣으세요!")
    
    student_file = st.file_uploader("📂 학생 생기부 파일 (PDF) 업로드", type=["pdf"], key="student_upload")
    st.markdown("**-- 또는 --**")
    student_text_input = st.text_area("📝 생기부 내용 직접 붙여넣기 (추천)", height=250)

with col2:
    st.subheader("2. 분석 옵션 및 추가 데이터 입력")
    teacher_context = st.text_area(
        "💡 특이사항 및 희망 전공 (예: 생명공학과 진학 희망)", 
        height=70
    )
    
    st.markdown("**🎯 목표 대학 전형 / 전공 가이드북 (선택)**")
    st.info("해당 대학의 가이드북을 업로드하면 평가 기준을 벤치마킹합니다. (※ 결과물에 특정 대학명은 노출되지 않습니다.)")
    univ_guide_file = st.file_uploader("🏫 대학 가이드북 PDF 업로드", type=["pdf"], key="univ_guide_upload")
    
    st.markdown("**📚 맞춤형 추천 도서 참고 자료 (선택)**")
    default_url = "https://nojaesu.com/category/DIRECTORY/%EA%B5%90%EA%B3%BC%EC%97%B0%EA%B3%84%26%EC%A0%84%EA%B3%B5%EC%A0%81%ED%95%A9%EC%84%9C%20%EA%B8%B0%EC%82%AC%20%EB%AA%A8%EC%9D%8C"
    book_url = st.text_input("🌐 추천 도서 웹사이트 주소(URL)", value=default_url)
    
    submit_btn = st.button("↵ 🚀 심층 분석 시작", type="primary", use_container_width=True)

st.markdown("---")

def create_word_file(text):
    doc = Document()
    doc.add_heading('AI 생기부 분석 결과 보고서', 0)
    doc.add_paragraph(text)
    
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

if submit_btn:
    if not api_key:
        st.error("왼쪽에 API 키를 먼저 입력해 주세요!")
    elif not pdf_files:
        st.error("기준이 될 PDF 파일이 폴더에 없습니다!")
    elif not student_file and not student_text_input.strip():
        st.error("학생의 생기부 파일(PDF)을 업로드하거나 텍스트를 직접 붙여넣어 주세요!")
    else:
        status_box = st.empty()
        
        try:
            status_box.info("⏳ [진행상황 1/5] 내장된 기본 가이드북을 학습하는 중입니다...")
            reference_text = load_reference_pdfs(pdf_files)
            
            univ_guide_text = ""
            if univ_guide_file:
                status_box.info("🏫 [진행상황 2/5] 업로드된 목표 대학 가이드북의 평가 기준을 분석 중입니다...")
                univ_pdf_reader = PyPDF2.PdfReader(univ_guide_file)
                for page in univ_pdf_reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        univ_guide_text += extracted + "\n"
            else:
                status_box.info("🏫 [진행상황 2/5] 목표 대학 가이드북이 생략되었습니다. 기본 범용 기준으로 진행합니다.")

            status_box.info("⏳ [진행상황 3/5] 학생의 생기부 데이터를 추출하는 중입니다...")
            student_data_text = ""
            
            if student_text_input.strip():
                student_data_text = student_text_input
            elif student_file:
                student_pdf_reader = PyPDF2.PdfReader(student_file)
                for page in student_pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        student_data_text += text + "\n"
            
            if not student_data_text.strip():
                raise Exception("생기부에서 글씨를 읽을 수 없습니다! PDF 대신 빈칸에 직접 붙여넣어 주세요.")
            
            status_box.info("📚 [진행상황 4/5] 추천 도서 목록 및 상세 본문을 수집하는 중입니다...")
            actual_book_data = ""
            
            if book_url.strip():
                try:
                    headers = {'User-Agent': 'Mozilla/5.0'}
                    response = requests.get(book_url.strip(), headers=headers)
                    response.raise_for_status() 
                    soup = BeautifulSoup(response.text, 'html.parser')
                    actual_book_data += soup.get_text(separator=' ', strip=True) + "\n\n"
                    
                    base_url = "/".join(book_url.split("/")[:3])
                    article_links = []
                    for a in soup.find_all('a', href=True):
                        href = a['href']
                        if re.match(r'^/[0-9]+(\?.*)?$', href) or "/entry/" in href:
                            full_url = base_url + href.split('?')[0]
                            if full_url not in article_links:
                                article_links.append(full_url)
                    
                    if article_links:
                        status_box.info(f"📚 [도서 연동] {len(article_links[:10])}개의 구체적인 도서 상세 설명을 추가로 수집 중입니다...")
                        for link in article_links[:10]:
                            try:
                                sub_res = requests.get(link, headers=headers, timeout=5)
                                sub_soup = BeautifulSoup(sub_res.text, 'html.parser')
                                content_area = sub_soup.find('div', class_='entry-content') or sub_soup.find('div', class_='article_view') or sub_soup.body
                                if content_area:
                                    actual_book_data += content_area.get_text(separator=' ', strip=True) + "\n\n"
                            except:
                                pass
                except Exception as e:
                    st.warning(f"⚠️ 입력하신 링크에 접속할 수 없습니다. (오류 메시지: {e})")
            
            if actual_book_data:
                actual_book_data = actual_book_data.replace("'쌤과 함께! 교과 연계 적합書]", "")
                actual_book_data = actual_book_data.replace("쌤과 함께! 교과 연계 적합書", "")
                actual_book_data = re.sub(r'[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳]', '', actual_book_data)

            book_instruction = ""
            if actual_book_data.strip():
                book_instruction = "반드시 아래 제공된 [추천 도서 참고 자료]의 텍스트 안에 '실제로 존재하는 책 제목과 저자'만 추출해서 추천하세요. 자료 안에 적합한 책이 없다면 억지로 지어내지 마세요."
            else:
                book_instruction = "별도로 제공된 도서 목록이 없으므로, AI가 자체적으로 학습한 실존하는 전공 적합 우수 도서를 추천해 주세요. (할루시네이션 절대 금지)"

            status_box.warning(f"🔍 [마무리 준비] {api_provider} API를 통해 심층 분석을 시작합니다...")
            
            # --- 💡 [프롬프트 대규모 고도화] 2번 약점 보완 강제 및 4번 도서 연계활동 양식 지정 ---
            prompt = f"""
            당신은 20년 경력의 대한민국 최고 수석 진학 상담 교사이자 입학사정관입니다.

            [담당 교사의 특별 지시사항 및 희망 전공]
            {teacher_context if teacher_context else "특별한 지시사항 없음."}
            
            [추천 도서 참고 자료 (게시물 상세 본문 포함)]
            {actual_book_data if actual_book_data else "제공된 목록 없음."}

            [기본 범용 대학 평가 기준 자료]
            {reference_text}

            [목표 대학 전형/전공 가이드북 평가 기준 (선택 사항)]
            {univ_guide_text if univ_guide_text else "제공된 목표 대학 가이드북 없음. 기본 범용 평가 기준만 적용할 것."}

            [업로드된 학생의 생기부 내용 (100% 팩트)]
            {student_data_text}

            🚨 [분석의 깊이 및 톤앤매너 (매우 중요)] 🚨
            1. 객관적이고 현실적인 평가 (과장 금지): 학생의 역량을 지나치게 긍정적으로 포장하거나 미사여구를 남발하는 것을 엄격히 금지합니다. 실제 입학사정관의 시각에서 사실 기반으로 건조하고 현실적으로 평가하세요.
            2. 단순 요약 금지: 단순히 생기부 내용을 요약하지 말고 '어떤 의미를 가지는지 깊이 있게 분석'하세요.

            🚨 [절대 엄수 - 출력 형식 및 규칙!] 🚨
            1. 🚫 [진행 중인 학년 기록 부재 지적 절대 금지!]: 최신 학년의 기록이 없는 것은 당연하므로 이를 단점으로 지적하지 마세요.
            2. 🚫 [출처 100% 일치 및 무단 나열 금지]: 문단 맨 앞에 '■ 소제목 [출처1, 출처2]' 형태로 확정한 출처와 본문 문장 끝의 개별 꼬리표는 100% 일치해야 합니다. 무단으로 나열하거나 즉흥적으로 빼먹지 마세요.
            3. 특정 대학명 노출 절대 금지: 동국대학교 등 대학 이름이 직접 등장하면 안 됩니다.
            4. 개조식 어미 사용: 문장 끝은 '~함', '~임', '~됨', '~판단됨' 으로 명사형 종결할 것.

            🚨 [항목별 세부 작성 규칙 (반드시 지킬 것!)] 🚨
            - 2번 항목(약점 분석): 반드시 최소 2가지 이상의 약점을 분석하세요. 약점을 지적할 때는 **반드시 학생부에 기재된 특정 활동 내용을 구체적으로 선언급(출처 꼬리표 포함)한 뒤**, 해당 활동의 한계를 짚고, 이를 보완하기 위한 구체적인 심화 탐구 방안이나 실질적 프로젝트를 제안하세요.
            - 4번 항목(추천 도서 및 연계 활동):
              1) 도서를 나열할 때 반드시 **"■ 교과명(또는 영역명):"** 이라는 굵은 소제목으로 과목/영역을 확실히 구분하세요.
              2) 각 과목 아래에 도서를 1번부터 차례대로 순차 번호를 매겨 나열하세요.
              3) 도서 1개당 아래의 2줄 포맷을 무조건 지키세요.
                 [번호]. <실제 책제목> (저자명 저) - 세부 본문 내용을 바탕으로 한 구체적인 도서 소개 및 추천 이유
                 **연계 활동:** 학생의 [O학년 OO과목] 내용과 매칭하여 구체적으로 어떤 심화 탐구/보고서 작성을 할 수 있는지 제안함.

            💡 [형식 참고용 정답 템플릿]
            ### 2. 범용 평가 기준에 비추어 볼 때 보완이 필요한 약점
            ■ 전공 심화 탐구 활동의 구체적인 실행 경험 부족 [1학년 진로활동, 1학년 통합과학]
            노화 원인 탐구에 대한 질문 확장을 통해 생명공학 분야에 깊은 관심은 드러냈으나 [1학년 진로활동], 유전자 가위 기술 탐구 등 단발적인 학습 경험에 그치고 본인만의 가설 설정과 검증 과정을 포함하는 확장된 탐구 경험으로 이어졌다는 내용은 확인하기 어려움 [1학년 통합과학]. 이를 보완하기 위해 가설을 세우고 실제 데이터를 수집, 분석하는 장기적인 심화 탐구 프로젝트를 기획하여 실천적 역량을 보여줄 필요가 있음.

            ### 4. 맞춤형 추천 도서 및 연계 활동 제안
            ■ 수학:
            1. <수학이 일상에서 이렇게 쓸모 있을 줄이야> (클라라 그리마 저) - 수학의 본질과 현실 속 쓰임을 풍부한 사례로 설명하여 학생의 호기심 해소에 도움이 됨.
            **연계 활동:** 이 책을 통해 얻은 아이디어를 바탕으로 [1학년 공통수학]에서 배운 이차함수 개념이 실생활에 적용되는 사례를 탐구하는 보고서를 작성함.
            
            ■ 생명과학:
            2. <세포처럼 나이 들 수 있다면> (김영웅 저) - 노화가 단일 원인이 아닌 복합적 과정임을 이해하고 텔로미어 마모에 대한 심화 질문을 확장하는 데 적합함.
            **연계 활동:** 독서 후 [1학년 진로활동]의 노화 탐구 기록과 연계하여 '노화 방지 기술의 윤리적 책임'을 주제로 논평을 작성하는 활동을 제안함.

            위의 모든 규칙과 템플릿을 완벽히 적용하여, 아래 5가지 양식에 맞추어 최종 결과물을 작성해 주세요.
            ### 1. 전공 적합성 및 주요 경쟁력 (테마별 엄선, 평가적 분석 서술, 100% 일치하는 분리 출처, 개조식)
            ### 2. 범용 평가 기준에 비추어 볼 때 보완이 필요한 약점 (엄선된 약점 분석, 100% 일치하는 분리 출처, 개조식) ※ 부재중인 학년의 기록 부족 지적 불가!
            ### 3. 추천 심화 탐구 주제 및 면접 예상 질문 3가지
            ### 4. 맞춤형 추천 도서 및 연계 활동 제안 (과목명 분류, 도서명(저자) 형식, **연계 활동:** 서술 필수)
            ### 5. 종합 의견 및 향후 발전 방향
            """
            
            result_text = ""
            
            # --- 💡 [신규] API 제공자에 따른 분기 처리 ---
            if api_provider == "Google AI Studio":
                best_model_name = ""
                for m in genai.list_models():
                    if 'generateContent' in m.supported_generation_methods:
                        best_model_name = m.name.replace("models/", "")
                        if 'flash' in best_model_name or 'pro' in best_model_name:
                            break 
                if best_model_name == "":
                    raise Exception("Google AI Studio에서 사용할 수 있는 모델이 없습니다.")
                model = genai.GenerativeModel(best_model_name)
                response = model.generate_content(prompt)
                result_text = response.text

            elif api_provider == "OpenRouter":
                headers = {
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json"
                }
                data = {
                    "model": or_model,
                    "messages": [{"role": "user", "content": prompt}]
                }
                res = requests.post("https://openrouter.ai/api/v1/chat/completions", headers=headers, json=data)
                res.raise_for_status() # 오류 발생 시 멈춤
                result_text = res.json()['choices'][0]['message']['content']

            status_box.success("✅ [분석 완료!] 심층 분석이 완료되었습니다. 결과물을 확인해 주세요!")
            st.write(result_text)
            
            word_file = create_word_file(result_text)
            st.download_button(
                label="📥 분석 결과 워드 다운로드",
                data=word_file,
                file_name="생기부_맞춤형_분석결과.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        except Exception as e:
            status_box.error(f"오류가 발생했습니다: {e}")

st.divider()
st.markdown("""
<div style='text-align: center; color: gray; padding: 20px; font-size: 13px;'>
    🏫 학교생활기록부 분석 시스템 v8.1<br>
    만든이: <b>신선여자고등학교 김명남</b>
</div>
""", unsafe_allow_html=True)
